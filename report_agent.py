from __future__ import annotations

import json
import os
from datetime import datetime
from typing import Any, Dict, Mapping, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


class ScorecardEvaluator:
    """Utility for computing a simple scorecard valuation."""

    CRITERIA: Dict[str, Mapping[str, Any]] = {
        "management_team": {"weight": 0.25, "label": "Management Team"},
        "market_opportunity": {"weight": 0.20, "label": "Market Opportunity"},
        "technology_product": {"weight": 0.18, "label": "Technology / Product"},
        "marketing_sales": {"weight": 0.15, "label": "Marketing & Sales"},
        "competitive_environment": {"weight": 0.10, "label": "Competitive Landscape"},
        "additional_funding": {"weight": 0.10, "label": "Funding Needs"},
        "other_factors": {"weight": 0.02, "label": "Other Factors"},
    }

    @classmethod
    def calculate(
        cls, scores: Mapping[str, Any], base_valuation: float = 7_000_000.0
    ) -> Dict[str, Any]:
        weighted_sum = 0.0
        details = []

        for key, meta in cls.CRITERIA.items():
            score = float(scores.get(key, 1.0))
            weighted = score * meta["weight"]
            weighted_sum += weighted
            details.append(
                {
                    "key": key,
                    "label": meta["label"],
                    "score": score,
                    "weight": meta["weight"],
                    "weighted": weighted,
                }
            )

        adjusted = base_valuation * weighted_sum

        return {
            "base_valuation": base_valuation,
            "weighted_sum": weighted_sum,
            "adjusted_valuation": adjusted,
            "details": details,
        }


def _coerce_text_field(value: Any) -> str:
    """Return the most useful string representation for downstream rendering."""
    if isinstance(value, str):
        return value
    if isinstance(value, Mapping):
        for key in ("full_text", "summary", "tech_summary", "market_summary"):
            text = value.get(key)
            if isinstance(text, str) and text.strip():
                return text
        return json.dumps(value, ensure_ascii=True, indent=2)
    if value is None:
        return ""
    return str(value)


def _safe_filename(name: str) -> str:
    filtered = "".join(ch for ch in name if ch.isalnum() or ch in ("-", "_"))
    return filtered or "startup"


def report_generator_agent(state: Dict[str, Any]) -> Dict[str, Any]:
    """Generate an investment memo in DOCX format based on the committee decision."""
    invest = int(state.get("investment_decision", 0)) == 1
    startup_name = state.get("startup_name", "Unknown Startup")
    startup_info = state.get("startup_info") or {}
    tech_summary = _coerce_text_field(state.get("tech_summary"))
    market_analysis = _coerce_text_field(state.get("market_analysis"))
    decision_reason = _coerce_text_field(state.get("decision_reason"))
    attempts = int(state.get("iteration_count", 0) or 0)
    reason_tokens = [token.strip() for token in decision_reason.split("|") if token.strip()]
    rationale_bullets: List[str] = []
    for token in reason_tokens:
        normalized = token.strip()
        lower = normalized.lower()
        simple = lower.replace(" ", "")
        if simple in {"투자", "비투자", "invest", "decline"}:
            continue
        if any(lower.startswith(prefix) for prefix in ("근거", "사유", "이유", "rationale", "reason")):
            payload = token.split(":", 1)[1] if ":" in token else token
            rationale_bullets.extend(seg.strip() for seg in payload.split(";") if seg.strip())
            continue
        rationale_bullets.append(normalized)

    invest_bullets: List[str] = []
    decline_bullets: List[str] = []
    if invest:
        invest_bullets = [bullet for bullet in rationale_bullets if bullet]
        if attempts > 1:
            invest_bullets.append(f"Decision reached after {attempts} validation loop(s), reflecting additional cross-checks before approval.")
        if not invest_bullets and decision_reason:
            invest_bullets.append(decision_reason)
        if not invest_bullets:
            invest_bullets.append("No explicit investment rationale recorded; review agent transcripts for additional detail.")
    else:
        decline_bullets = [bullet for bullet in rationale_bullets if bullet]
        if not decline_bullets and decision_reason:
            decline_bullets.append(decision_reason)
        if attempts >= 5:
            decline_bullets.append(f"{attempts} consecutive evaluation loops exhausted the retry budget without meeting evidence coverage thresholds.")
        elif attempts > 1:
            decline_bullets.append(f"Decision was reached after {attempts} iterations with insufficient conviction.")
        if not decline_bullets:
            decline_bullets.append("No explicit decline rationale recorded; inspect LangGraph logs for contributing factors.")

    scorecard_scores = state.get("scorecard_scores") or {}
    valuation = ScorecardEvaluator.calculate(scorecard_scores)

    doc = Document()

    # Cover
    title = doc.add_heading(startup_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("AI Investment Memo", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    method_para = doc.add_paragraph("Produced via Scorecard Valuation Method.")
    method_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    method_run = method_para.runs[0]
    method_run.font.size = Pt(12)
    method_run.italic = True

    doc.add_paragraph()
    date_para = doc.add_paragraph(datetime.now().strftime("%Y-%m-%d"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")
    banner = doc.add_paragraph("INVESTMENT RECOMMENDED" if invest else "INVESTMENT DECLINED")
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.runs[0]
    banner_run.font.size = Pt(22)
    banner_run.font.bold = True
    banner_run.font.color.rgb = RGBColor(0, 112, 192) if invest else RGBColor(192, 0, 0)

    doc.add_page_break()

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    if invest:
        summary_intro = (
            f"{startup_name} operates in the {startup_info.get('tech_focus', 'N/A')} domain. "
            f"Based on the scorecard assessment, the adjusted pre-money valuation is "
            f"${valuation['adjusted_valuation']:,.0f}. The investment committee "
            "recommends proceeding subject to confirmatory diligence."
        )
    else:
        summary_intro = (
            f"{startup_name} operates in the {startup_info.get('tech_focus', 'N/A')} domain. "
            f"The adjusted pre-money valuation benchmarked via scorecard is "
            f"${valuation['adjusted_valuation']:,.0f}, yet the committee recommends "
            "declining the opportunity at this time given outstanding risks."
        )
    doc.add_paragraph(summary_intro)

    if invest:
        if invest_bullets:
            doc.add_paragraph("Key investment highlights:")
            for bullet in invest_bullets[:3]:
                doc.add_paragraph(bullet, style="List Bullet")
        if attempts > 1:
            doc.add_paragraph(
                f"The committee completed {attempts} validation loop(s) before approving the deal."
            )
    else:
        if attempts:
            doc.add_paragraph(
                f"The LangGraph workflow completed {attempts} iteration(s) without reaching the required evidence coverage thresholds."
            )
        if decline_bullets:
            doc.add_paragraph("Key decline drivers (condensed):")
            for bullet in decline_bullets[:3]:
                doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph()
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = "Light Grid Accent 1"
    summary_points = [
        ("Company", startup_name),
        ("Focus Area", startup_info.get("tech_focus", "N/A")),
        ("Founded", startup_info.get("founded_year", "N/A")),
        ("Adjusted Valuation", f"${valuation['adjusted_valuation']:,.0f}"),
        ("Decision", "Invest" if invest else "Decline"),
    ]
    for idx, (label, value) in enumerate(summary_points):
        summary_table.rows[idx].cells[0].text = str(label)
        summary_table.rows[idx].cells[1].text = str(value)

    doc.add_page_break()

    # Sections
    doc.add_heading("1. Company Overview", level=1)
    overview_table = doc.add_table(rows=6, cols=2)
    overview_table.style = "Light Shading Accent 1"
    overview_rows = [
        ("Website", startup_info.get("website", "N/A")),
        ("Headquarters", startup_info.get("location", "N/A")),
        ("Funding", startup_info.get("funding", "N/A")),
        ("Team Size", startup_info.get("team_size", "N/A")),
        ("Business Description", startup_info.get("description", "N/A")),
        ("Additional Notes", startup_info.get("notes", "N/A")),
    ]
    for idx, (label, value) in enumerate(overview_rows):
        overview_table.rows[idx].cells[0].text = str(label)
        overview_table.rows[idx].cells[1].text = _coerce_text_field(value)

    doc.add_page_break()

    doc.add_heading("2. Technology & Team Assessment", level=1)
    doc.add_paragraph(tech_summary or "No technology summary provided.")

    doc.add_page_break()

    doc.add_heading("3. Market Assessment", level=1)
    doc.add_paragraph(market_analysis or "No market analysis provided.")

    doc.add_page_break()

    doc.add_heading("4. Scorecard Valuation", level=1)
    scoring_table = doc.add_table(rows=len(valuation["details"]) + 1, cols=4)
    scoring_table.style = "Light Grid Accent 1"
    headers = ["Criteria", "Score", "Weight", "Weighted"]
    for col, header in enumerate(headers):
        scoring_table.rows[0].cells[col].text = header

    for row_idx, detail in enumerate(valuation["details"], start=1):
        scoring_table.rows[row_idx].cells[0].text = detail["label"]
        scoring_table.rows[row_idx].cells[1].text = f"{detail['score']:.2f}"
        scoring_table.rows[row_idx].cells[2].text = f"{detail['weight']:.0%}"
        scoring_table.rows[row_idx].cells[3].text = f"{detail['weighted']:.4f}"

    doc.add_paragraph()
    doc.add_paragraph(
        "Score interpretations: values below 1.0 lag the benchmark, "
        "1.0 tracks the benchmark, and values above 1.0 outperform."
    )

    doc.add_page_break()

    doc.add_heading("5. Investment Considerations", level=1)
    doc.add_heading("Decision Rationale", level=2)
    if invest:
        if invest_bullets:
            doc.add_paragraph("Investment thesis and validation notes:")
            for bullet in invest_bullets:
                doc.add_paragraph(bullet, style="List Bullet")
        else:
            doc.add_paragraph(decision_reason or "No explicit rationale captured.")
    else:
        if attempts:
            doc.add_paragraph(
                f"The committee completed {attempts} evaluation loop(s) without gathering sufficient conviction to proceed."
            )
        if decline_bullets:
            doc.add_paragraph("Detailed decline rationale:")
            for bullet in decline_bullets:
                doc.add_paragraph(bullet, style="List Bullet")
        else:
            doc.add_paragraph(decision_reason or "No explicit rationale captured.")

    strengths = [
        f"{detail['label']}: {detail['score']:.2f}"
        for detail in valuation["details"]
        if detail["score"] >= 1.2
    ]
    doc.add_heading("Strengths", level=2)
    if strengths:
        for item in strengths:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No criteria exceeded the 1.2 threshold.")

    risks = [
        f"{detail['label']}: {detail['score']:.2f}"
        for detail in valuation["details"]
        if detail["score"] < 0.9
    ]
    doc.add_heading("Watchpoints", level=2)
    if risks:
        for item in risks:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No criteria fell below the 0.9 threshold.")

    doc.add_page_break()

    doc.add_heading("6. Recommendation & Next Steps", level=1)
    if invest:
        doc.add_paragraph(
            "Proceed with investment subject to confirmatory diligence on financials, "
            "customer references, and regulatory requirements."
        )
        doc.add_paragraph(
            f"Suggested cheque size: ${valuation['adjusted_valuation'] * 0.2:,.0f} "
            "(approximately 20% of the adjusted valuation)."
        )
    else:
        doc.add_paragraph(
            "Do not commit capital at this time. Revisit the opportunity if key risks are "
            "mitigated or new traction emerges."
        )
        doc.add_paragraph(
            "Recommended follow-ups: maintain light-touch monitoring, request quarterly updates, "
            "and re-open full diligence if core metrics improve materially."
        )

    output_dir = os.path.join(os.getcwd(), "results")
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{_safe_filename(startup_name)}_investment_report_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    state["report"] = output_path
    return state
