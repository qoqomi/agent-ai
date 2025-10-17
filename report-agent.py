from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


class ScorecardEvaluator:
    """Scorecard Valuation Method 평가 기준"""

    CRITERIA = {
        "management_team": {"weight": 0.25, "name": "경영진/팀 역량"},
        "market_opportunity": {"weight": 0.20, "name": "시장 기회 규모"},
        "technology_product": {"weight": 0.18, "name": "기술/제품 경쟁력"},
        "marketing_sales": {"weight": 0.15, "name": "마케팅/영업 역량"},
        "competitive_environment": {"weight": 0.10, "name": "경쟁 환경"},
        "additional_funding": {"weight": 0.10, "name": "추가 자금 필요성"},
        "other_factors": {"weight": 0.02, "name": "기타 요소"},
    }

    @classmethod
    def calculate_valuation(
        cls, scores: dict, base_valuation: float = 7_000_000
    ) -> dict:
        """
        Scorecard 방식으로 Pre-money Valuation 계산
        scores: 각 평가 항목별 점수 (0.0 ~ 2.0, 1.0이 평균)
        """
        weighted_sum = 0.0
        details = []

        for key, criteria in cls.CRITERIA.items():
            score = scores.get(key, 1.0)
            weighted_score = score * criteria["weight"]
            weighted_sum += weighted_score

            details.append(
                {
                    "criteria": criteria["name"],
                    "score": score,
                    "weight": criteria["weight"],
                    "weighted": weighted_score,
                }
            )

        adjusted_valuation = base_valuation * weighted_sum

        return {
            "base_valuation": base_valuation,
            "weighted_sum": weighted_sum,
            "adjusted_valuation": adjusted_valuation,
            "details": details,
        }


def report_generator_agent(state: dict) -> dict:
    """
    AI 스타트업 투자 평가 보고서 생성 에이전트
    Scorecard Valuation Method 기반 정량 평가 포함
    """

    investment_decision = state.get("investment_decision", 0)

    if investment_decision != 1:
        state["report"] = "투자 보류로 인해 보고서 미생성"
        return state

    # 데이터 추출
    company_name = state["startup_name"]
    company_info = state["startup_info"]
    tech_summary = state.get("tech_summary", "기술 정보 없음")
    market_analysis = state.get("market_analysis", "시장 정보 없음")
    decision_reason = state.get("decision_reason", "")

    # Scorecard 평가 점수 (실제 구현시 AI 분석 결과 활용)
    scorecard_scores = state.get(
        "scorecard_scores",
        {
            "management_team": 1.3,
            "market_opportunity": 1.4,
            "technology_product": 1.2,
            "marketing_sales": 0.9,
            "competitive_environment": 1.1,
            "additional_funding": 1.0,
            "other_factors": 1.0,
        },
    )

    # Scorecard 평가 계산
    valuation_result = ScorecardEvaluator.calculate_valuation(scorecard_scores)

    # 워드 문서 생성
    doc = Document()

    # ========================================
    # 표지
    # ========================================
    title = doc.add_heading(f"{company_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("AI 스타트업 투자 평가 보고서", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n" * 2)

    method_para = doc.add_paragraph("Scorecard Valuation Method 기반")
    method_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    method_para.runs[0].font.size = Pt(12)
    method_para.runs[0].italic = True

    doc.add_paragraph("\n" * 2)

    date_para = doc.add_paragraph(f"{datetime.now().strftime('%Y년 %m월 %d일')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n" * 3)

    # 투자 추천 표시
    recommend = doc.add_paragraph("투자 추천")
    recommend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    recommend.runs[0].font.size = Pt(24)
    recommend.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    recommend.runs[0].bold = True

    doc.add_page_break()

    # ========================================
    # Executive Summary
    # ========================================
    doc.add_heading("Executive Summary", 1)

    doc.add_paragraph(
        f"{company_name}는 AI 기술 기반 {company_info.get('tech_focus', 'N/A')} 분야의 혁신 스타트업입니다. "
        f"Scorecard Valuation Method를 활용한 정량 평가 결과, "
        f"기업 가치는 약 ${valuation_result['adjusted_valuation']:,.0f}로 산정되며, "
        f"투자 가치가 높은 것으로 판단됩니다."
    )

    doc.add_paragraph()

    # 핵심 지표
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = "Light Grid Accent 1"

    summary_data = [
        ("회사명", company_name),
        ("사업 분야", company_info.get("tech_focus", "N/A")),
        ("설립연도", str(company_info.get("founded_year", "N/A"))),
        ("평가 기업가치", f"${valuation_result['adjusted_valuation']:,.0f}"),
        ("투자 판단", "투자 추천"),
    ]

    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = str(value)

    doc.add_page_break()

    # ========================================
    # 목차
    # ========================================
    doc.add_heading("목차", 1)

    toc = [
        "1. 회사 개요",
        "2. 기술 분석",
        "3. 시장 분석",
        "4. Scorecard 정량 평가",
        "5. 투자자 관점 분석",
        "6. 투자 의견 및 권장사항",
        "7. 결론",
        "부록: 평가 방법론",
    ]

    for item in toc:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # ========================================
    # 1. 회사 개요
    # ========================================
    doc.add_heading("1. 회사 개요", 1)

    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = "Light Shading Accent 1"

    info_data = [
        ("회사명", company_name),
        ("웹사이트", company_info.get("website", "N/A")),
        ("사업 분야", company_info.get("tech_focus", "N/A")),
        ("설립연도", str(company_info.get("founded_year", "N/A"))),
        ("소재지", company_info.get("location", "N/A")),
        ("투자 유치 현황", company_info.get("funding", "N/A")),
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value

    doc.add_paragraph()
    doc.add_heading("사업 개요", 2)
    doc.add_paragraph(company_info.get("description", "정보 없음"))

    doc.add_page_break()

    # ========================================
    # 2. 기술 분석
    # ========================================
    doc.add_heading("2. 기술 분석", 1)
    doc.add_paragraph(tech_summary)

    doc.add_page_break()

    # ========================================
    # 3. 시장 분석
    # ========================================
    doc.add_heading("3. 시장 분석", 1)
    doc.add_paragraph(market_analysis)

    doc.add_page_break()

    # ========================================
    # 4. Scorecard 정량 평가
    # ========================================
    doc.add_heading("4. Scorecard 정량 평가", 1)

    doc.add_paragraph(
        "본 평가는 Angel Capital Association에서 제시한 Scorecard Valuation Method를 "
        "기반으로 수행되었습니다. 이 방법론은 초기 단계 스타트업의 가치를 평가하는 "
        "가장 널리 사용되는 정량적 접근법입니다."
    )

    doc.add_paragraph()

    # 평가 결과 요약
    doc.add_heading("평가 결과 요약", 2)

    result_table = doc.add_table(rows=3, cols=2)
    result_table.style = "Light Grid Accent 1"

    result_data = [
        ("기준 기업가치", f"${valuation_result['base_valuation']:,.0f}"),
        ("가중치 합계", f"{valuation_result['weighted_sum']:.4f}"),
        ("조정 기업가치", f"${valuation_result['adjusted_valuation']:,.0f}"),
    ]

    for i, (label, value) in enumerate(result_data):
        result_table.rows[i].cells[0].text = label
        result_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 세부 평가 항목
    doc.add_heading("세부 평가 항목", 2)

    detail_table = doc.add_table(rows=len(valuation_result["details"]) + 1, cols=4)
    detail_table.style = "Light Shading Accent 1"

    # 헤더
    headers = ["평가 항목", "점수", "가중치", "가중 점수"]
    for i, header in enumerate(headers):
        cell = detail_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # 데이터
    for i, detail in enumerate(valuation_result["details"], 1):
        detail_table.rows[i].cells[0].text = detail["criteria"]
        detail_table.rows[i].cells[1].text = f"{detail['score']:.2f}"
        detail_table.rows[i].cells[2].text = f"{detail['weight']:.0%}"
        detail_table.rows[i].cells[3].text = f"{detail['weighted']:.4f}"

    doc.add_paragraph()

    # 평가 기준 설명
    doc.add_heading("평가 기준 설명", 2)

    doc.add_paragraph(
        "• 점수 범위: 0.0 ~ 2.0 (1.0이 시장 평균)\n"
        "• 1.0 미만: 시장 평균 이하\n"
        "• 1.0: 시장 평균 수준\n"
        "• 1.0 초과: 시장 평균 이상\n"
        "• 1.5 이상: 매우 우수한 수준"
    )

    doc.add_page_break()

    # ========================================
    # 5. 투자자 관점 분석
    # ========================================
    doc.add_heading("5. 투자자 관점 분석", 1)

    doc.add_heading("강점 분석", 2)

    strengths = []
    for detail in valuation_result["details"]:
        if detail["score"] >= 1.2:
            strengths.append(f"• {detail['criteria']}: {detail['score']:.2f} (우수)")

    if strengths:
        for strength in strengths:
            doc.add_paragraph(strength)
    else:
        doc.add_paragraph("시장 평균 수준 유지")

    doc.add_paragraph()

    doc.add_heading("개선 영역", 2)

    weaknesses = []
    for detail in valuation_result["details"]:
        if detail["score"] < 1.0:
            weaknesses.append(
                f"• {detail['criteria']}: {detail['score']:.2f} (개선 필요)"
            )

    if weaknesses:
        for weakness in weaknesses:
            doc.add_paragraph(weakness)
    else:
        doc.add_paragraph("전반적으로 균형잡힌 역량 보유")

    doc.add_paragraph()

    if decision_reason:
        doc.add_heading("종합 의견", 2)
        doc.add_paragraph(decision_reason)

    doc.add_page_break()

    # ========================================
    # 6. 투자 의견 및 권장사항
    # ========================================
    doc.add_heading("6. 투자 의견 및 권장사항", 1)

    doc.add_heading("투자 추천", 2)
    recommend_para = doc.add_paragraph(
        f"Scorecard 평가 결과 가중치 합계 {valuation_result['weighted_sum']:.2f}로, "
        "투자를 추천합니다."
    )
    recommend_para.runs[0].font.size = Pt(12)
    recommend_para.runs[0].font.bold = True
    recommend_para.runs[0].font.color.rgb = RGBColor(0, 112, 192)

    doc.add_paragraph()

    doc.add_heading("투자 근거", 2)
    doc.add_paragraph(
        f"1. 정량 평가: Scorecard 방식 평가 결과 ${valuation_result['adjusted_valuation']:,.0f}의 "
        "기업 가치 산정\n\n"
        "2. AI 기술 경쟁력: 핵심 기술 및 차별화 요소 보유\n\n"
        "3. 시장 성장성: 타겟 시장의 높은 성장 잠재력\n\n"
        "4. 팀 역량: 경험 있는 경영진 및 기술팀 구성"
    )

    doc.add_paragraph()

    doc.add_heading("권장 사항", 2)
    doc.add_paragraph(
        f"• 투자 규모: ${valuation_result['adjusted_valuation'] * 0.15:,.0f} ~ "
        f"${valuation_result['adjusted_valuation'] * 0.25:,.0f} (기업가치의 15-25%)\n"
        "• 상세 실사(Due Diligence) 진행\n"
        "• 단계별 투자 계획 수립 (Seed → Series A)\n"
        "• 투자 후 정기 모니터링 체계 구축\n"
        "• 3-5년 Exit 전략 수립"
    )

    doc.add_page_break()

    # ========================================
    # 7. 결론
    # ========================================
    doc.add_heading("7. 결론", 1)

    doc.add_paragraph(
        f"{company_name}는 AI 기술 기반 {company_info.get('tech_focus', 'N/A')} 분야에서 "
        "경쟁력 있는 기술과 팀을 보유한 초기 단계 스타트업입니다."
    )

    doc.add_paragraph()

    doc.add_paragraph(
        f"Scorecard Valuation Method 기반 정량 평가 결과, "
        f"기업 가치는 약 ${valuation_result['adjusted_valuation']:,.0f}로 산정되며, "
        f"가중치 합계 {valuation_result['weighted_sum']:.2f}로 시장 평균을 상회하는 수준입니다."
    )

    doc.add_paragraph()

    doc.add_paragraph(
        "종합 평가 결과, 본 기업은 투자 가치가 충분하며 향후 성장 가능성이 높은 것으로 판단됩니다. "
        "상세 실사 진행 후 투자 규모 및 조건을 확정할 것을 제안합니다."
    )

    doc.add_paragraph("\n\n")

    # 서명
    signature = doc.add_paragraph(
        f"평가 완료일: {datetime.now().strftime('%Y년 %m월 %d일')}\n" "AI 투자심사팀"
    )
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ========================================
    # 부록: 평가 방법론
    # ========================================
    doc.add_page_break()
    doc.add_heading("부록: 평가 방법론", 1)

    doc.add_heading("Scorecard Valuation Method", 2)
    doc.add_paragraph(
        "Scorecard Valuation Method는 Bill Payne이 개발한 초기 단계 스타트업 가치평가 방법론으로, "
        "Angel Capital Association에서 공식 인정하는 표준 평가 방식입니다."
    )

    doc.add_paragraph()

    doc.add_heading("평가 프로세스", 2)
    doc.add_paragraph(
        "1. 동일 지역, 동일 업종의 유사 스타트업 Pre-money Valuation 중간값 산정\n"
        "2. 7개 핵심 평가 요소별 점수 산정 (0.0 ~ 2.0)\n"
        "3. 각 요소별 가중치 적용하여 가중 점수 계산\n"
        "4. 가중 점수 합계를 기준 기업가치에 곱하여 최종 기업가치 도출"
    )

    doc.add_paragraph()

    doc.add_heading("참고 문헌", 2)
    doc.add_paragraph(
        "• Bill Payne, Scorecard Valuation Methodology\n"
        "• Angel Capital Association (2019), Angel Funders Report\n"
        "• Reference: https://angelcapitalassociation.org"
    )

    # ========================================
    # 파일 저장
    # ========================================
    result_dir = os.path.join(os.getcwd(), "results")
    os.makedirs(result_dir, exist_ok=True)

    safe_name = company_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
    filename = f"AI_Investment_Report_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(result_dir, filename)

    doc.save(filepath)
    state["report"] = filepath

    print(f"✅ 투자 평가 보고서 생성 완료: {filename}")

    return state


# 사용 예시
if __name__ == "__main__":
    # 테스트 데이터
    # test_state = {
    #     "investment_decision": 1,
    #     "startup_name": "AI Vision Labs",
    #     "startup_info": {
    #         "website": "www.aivisionlabs.com",
    #         "tech_focus": "컴퓨터 비전 및 자율주행",
    #         "founded_year": 2023,
    #         "location": "서울, 대한민국",
    #         "funding": "Pre-Seed $500K",
    #         "description": "AI 기반 실시간 객체 인식 및 추적 기술을 활용한 자율주행 솔루션 개발",
    #     },
    #     "tech_summary": "자체 개발한 경량화 딥러닝 모델로 엣지 디바이스에서 실시간 처리 가능",
    #     "market_analysis": "글로벌 자율주행 시장은 연평균 38% 성장 예상",
    #     "decision_reason": "우수한 기술팀과 차별화된 AI 기술력 보유",
    #     "scorecard_scores": {
    #         "management_team": 1.3,
    #         "market_opportunity": 1.4,
    #         "technology_product": 1.25,
    #         "marketing_sales": 0.9,
    #         "competitive_environment": 1.1,
    #         "additional_funding": 1.0,
    #         "other_factors": 1.05,
    #     },
    # }

    test_state = {
        "startup_name": "OrbitalMesh",
        "startup_info": {
            "is_startup": true,
            "weekly_growth": 8.1,
            "cash_on_hand": 1500000,
            "monthly_burn": 70000,
            "ltv": 2400,
            "cac": 700,
        },
        "tech_summary": {
            "full_text": "기술 평가\n핵심 기술: 초소형 큐브샛을 위한 지상국 시스템 엔지니어링 솔루션 및 AI 기반 위성영상 분석 기술.  \n개발 단계: 상용화  \n차별화: \n1. AI 딥러닝 기반의 위성영상 품질 보정 기술 제공.\n2. 지상국 안테나 및 모뎀 제작으로 통합 솔루션 제공.\n3. UAE를 포함한 해외 시장으로의 위성통신 단말기 공급.  \n강점: \n1. 8년 만에 코스닥 상장으로 시장 신뢰도 확보.\n2. 다양한 위성 관련 제품군 보유.\n3. 우주산업 솔루션에 대한 전문성.  \n약점: \n1. 초기 자금 조달의 어려움 경험.\n2. 경쟁이 치열한 우주산업 내에서의 시장 점유율 확보 필요.  \n기술 점수: 120%\n\n경쟁사 비교\n| 항목 | 컨텍 | 경쟁사A | 경쟁사B |\n|------|------|----------|----------|\n| 크기 | 중소기업 | 대기업 | 중소기업 |\n| 성능 | AI 기반 분석 우수 | 전통적 분석 | 제한적 기능 |\n\n팀 평가\n창업자: CEO 이성희 (우주산업 경력 10년 이상), CTO [이름/경력 미제공]  \n팀 규모: 50명  \n핵심 역량: \n1. 우주산업 솔루션 개발 경험.\n2. AI 및 데이터 분석 기술.\n3. 지상국 시스템 엔지니어링 전문성.  \n산업 경험: 있음 (우주산업 및 관련 기술 분야에서의 경험)  \n팀 점수: 110%\n\n종합 리스크\n- 초기 자금 조달의 어려움으로 인한 재무적 불안정성.\n- 치열한 경쟁 속에서의 시장 점유율 확보의 어려움.",
            "tech_score": 120,
            "team_score": 110,
            "tech_summary": "초소형 큐브샛을 위한 지상국 시스템 엔지니어링 솔루션 및 AI 기반 위성영상 분석 기술.",
            "team_summary": "창업자: CEO 이성희 (우주산업 경력 10년 이상), CTO [이름/경력 미제공] 팀 규모: 50명 핵심 역량:",
            "key_risks": ["리스크 정보 없음"],
            "competitive_table": "| 항목 | 컨텍 | 경쟁사A | 경쟁사B |\n|------|------|----------|----------|\n| 크기 | 중소기업 | 대기업 | 중소기업 |\n| 성능 | AI 기반 분석 우수 | 전통적 분석 | 제한적 기능 |",
            "sources": [
                "https://brunch.co.kr/@cliche-cliche/106",
                "https://m.blog.naver.com/ivvlove/222445956219",
                "https://www.jake-james.com/blog/what-does-the-ceo-cfo-coo-cto-and-others-do",
            ],
        },
        "market_analysis": "해양/원격 산업 IoT를 타깃. Why Now: 발사 단가 하락·오지 연결 수요 증가.",
        "investment_decision": 1,
        "decision_reason": "결정: 투자 | 기술 120%, 시장 130% | 근거: 기술적으로 AI 기반 위성영상 분석 및 지상국 시스템 엔지니어링 솔루션을 제공하여 차별화된 경쟁력을 보유하고 있음.; 해양/원격 산업 IoT를 타겟으로 하여 발사 단가 하락과 오지 연결 수요 증가로 시장 기회가 확대되고 있음.; 팀의 경험과 전문성이 우주산업에 적합하며, 시장 신뢰도를 확보한 상태임.# 투자 판단 — OrbitalMesh - **결과**: **투자** (label=1) - **Scorecard(간이)**: 기술 120% · 시장 130% ## 판단 근거(LLM 요약) - 기술적으로 AI 기반 위성영상 분석 및 지상국 시스템 엔지니어링 솔루션을 제공하여 차별화된 경쟁력을 보유하고 있음. - 해양/원격 산업 IoT를 타겟으로 하여 발사 단가 하락과 오지 연결 수요 증가로 시장 기회가 확대되고 있음. - 팀의 경험과 전문성이 우주산업에 적합하며, 시장 신뢰도를 확보한 상태임. ## RAG 인용(상위 3 / 각 피벗) **PMF** - Product/market fit means being in a good market with a product that can satisfy that market. You can always feel when product/market fit isn't happening. The customers aren't quite getting value out of the product, word of mouth is n't spreading, usage isn't growing that fast, pr … (Unknown·PMF·p.25) - And you can always feel product/market fit when it's happening. The customers are buying the product just as fast as you can make it --or usage is growing just as fast as you can add more servers. Money from customers is piling up in your company checking account. You're hiring s … (Unknown·PMF·p.26) **Traction** - You look at a startup and ask, will this team be able to optimally execute again st their opportunity? I focus on effectiveness as opposed to experience, since the history of the tech industry is full of highly successful startups that were staffed primarily by people who had nev … (Unknown·Traction·p.21) - The other connection between startups and technology is that startups create new ways of doing things, and new ways of doing things are, in the broader sense of the word, new technology. When a start up both begins with an idea exposed by technological change and makes a product … (Unknown·Traction·p.3) - By: Bill Payne, Frontier Angels This article was originally written in May 2001 and has been updated multiple times. Others have referred to this and similar methods as the Benchmark Method and the Bill Payne Method. The Scorecard Valuation Methodology is useful for investment in … (Bill Gurley·Traction·p.14) **DefaultAlive** - Why do so few founders know whether they're default alive or default dead? Mainly, I think, because they're not used to asking that. It's not a question that makes sense to ask e arly on, any more than it makes sense to ask a 3 year old how he plans to support himself. But as the … (Unknown·DefaultAlive·p.28) **Team** - product and intellectual property are important, but the quality of the team is key. Making the Valuation Calculation To provide an example, assume a co mpany with an average product and technology (100% of norm), a strong team (125% of norm) and a large market opportunity (150% … (Unknown·Team·p.16) - by how appealing it is to lots of customers? No. Product quality and market size are completely different. Here's the classic scenario: the world's best software application for an operating system nobody runs. Just ask any software developer targeting the market for BeOS, Amiga, … (Unknown·Team·p.21) - Strength of Entrepreneur and Team 30% max 125% 0.3750 Size of the Opportunity 25% max 150% 0.3750 Product/Technology 15% max 100% 0.1500 Competitive Environment 10% max 75% 0.0750 Marketing/Sales/Partnerships 10% max 80% 0.0800 Need for Additional Investment 5% max 100% 0.0500 Ot … (Unknown·Team·p.17)",
        "report": "",
        "iteration_count": 1,
    }

    result = report_generator_agent(test_state)
    print(f"\n보고서 경로: {result['report']}")
